import streamlit as st
from tavily import TavilyClient
from dotenv import load_dotenv
import os
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from io import BytesIO


# Load API keys
load_dotenv()
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

client = TavilyClient(TAVILY_API_KEY)

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")

# st.set_page_config(page_title="Deep Research Assistant", layout="wide")
# st.title("Deep Research Assistant")

# # Topic selection outside the form so it updates immediately
# topic = st.selectbox("Search Topic", ["general", "news", "finance"])

# # Form using updated topic
# with st.form("search_form"):
#     col1, col2 = st.columns(2)
#     with col1:
#         query = st.text_area("Query", "hello", height=200)
#         search_depth = st.selectbox("Search Depth", ["basic", "advanced"])
#         max_results = st.slider("Select Depth (max=10)", 1, 10, 5)
        
        
#     with col2:
#         include_answer = st.selectbox("Include Answer", ["none", "basic", "advanced"])
#         include_images = st.checkbox("Include Images")
#         country = st.selectbox(
#             "Country (only for general)",
#             [ "none", "afghanistan", "albania", "algeria", "andorra", "angola", "argentina", "armenia", "australia", "austria", "azerbaijan", "bahamas", "bahrain", "bangladesh", "barbados", "belarus", "belgium", "belize", "benin", "bhutan", "bolivia", "bosnia and herzegovina", "botswana", "brazil", "brunei", "bulgaria", "burkina faso", "burundi", "cambodia", "cameroon", "canada", "cape verde", "central african republic", "chad", "chile", "china", "colombia", "comoros", "congo", "costa rica", "croatia", "cuba", "cyprus", "czech republic", "denmark", "djibouti", "dominican republic", "ecuador", "egypt", "el salvador", "equatorial guinea", "eritrea", "estonia", "ethiopia", "fiji", "finland", "france", "gabon", "gambia", "georgia", "germany", "ghana", "greece", "guatemala", "guinea", "haiti", "honduras", "hungary", "iceland", "india", "indonesia", "iran", "iraq", "ireland", "israel", "italy", "jamaica", "japan", "jordan", "kazakhstan", "kenya", "kuwait", "kyrgyzstan", "latvia", "lebanon", "lesotho", "liberia", "libya", "liechtenstein", "lithuania", "luxembourg", "madagascar", "malawi", "malaysia", "maldives", "mali", "malta", "mauritania", "mauritius", "mexico", "moldova", "monaco", "mongolia", "montenegro", "morocco", "mozambique", "myanmar", "namibia", "nepal", "netherlands", "new zealand", "nicaragua", "niger", "nigeria", "north korea", "north macedonia", "norway", "oman", "pakistan", "panama", "papua new guinea", "paraguay", "peru", "philippines", "poland", "portugal", "qatar", "romania", "russia", "rwanda", "saudi arabia", "senegal", "serbia", "singapore", "slovakia", "slovenia", "somalia", "south africa", "south korea", "south sudan", "spain", "sri lanka", "sudan", "sweden", "switzerland", "syria", "taiwan", "tajikistan", "tanzania", "thailand", "togo", "trinidad and tobago", "tunisia", "turkey", "turkmenistan", "uganda", "ukraine", "united arab emirates", "united kingdom", "united states", "uruguay", "uzbekistan", "venezuela", "vietnam", "yemen", "zambia", "zimbabwe"],
#             disabled=(topic != "general")
#         )
#         days = st.slider("Days (only for news)", 1, 7, 6, disabled=(topic != "news"))
#         time_range = st.selectbox("Time Range", ["none", "day", "week", "month", "year"])

#     submitted = st.form_submit_button("Search")

# if submitted:
#     with st.spinner("Searching with Tavily..."):
#         search_params = {
#             "query": query,
#             "search_depth": search_depth,
#             "max_results": max_results,
#             "time_range": time_range,
#             "include_answer": include_answer,
#             "include_images": include_images,
#         }
        # if time_range == "none":
        #     del search_params["time_range"]
        # if include_answer == "none":
        #     search_params["include_answer"] = False
        # if topic == "news":
        #     search_params["days"] = days
        # if topic == "general":
        #     search_params["country"] = country

#         response = client.search(**search_params)

#         st.success("Search completed!")
#         st.json(response)  # Show raw response

#         content_blocks = []
#         for result in response["results"]:
#             text = result.get("content", "")
#             images = result.get("images", [])
#             content_blocks.append(text)
#             for img_url in images:
#                 content_blocks.append(f"[IMAGE]({img_url})")

#         full_content = "\n\n".join(content_blocks)

#         # Gemini Prompt
#         prompt = f"""
# You are an expert research analyst and report writer. Your task is to create a deep, detailed, structured, and insightful research report on the following topic:

# **Topic:** {query}

# Use the provided content below as your primary source of information. Analyze it thoroughly. Extract facts, insights, and diverse perspectives. Your report should follow this structure:

# ---

# ### 1. Executive Summary
# Summarize the key findings and importance of the topic in under 150 words.

# ### 2. Introduction
# Explain the background and context. Why is it relevant now?

# ### 3. Key Highlights
# - Summarize major findings or trends.
# - Use bullet points.

# ### 4. In-depth Analysis
# Break down the topic into multiple sections. For each:
# - Explain clearly.
# - Use quotes or paraphrased points from sources.
# - Mention supporting images where relevant.

# ### 5. Recent Developments (if any)
# Highlight newsworthy events.

# ### 6. Regional or Industry Insights
# (Optional based on content)

# ### 7. Conclusion
# Summarize and suggest future outlooks or actions.

# ### 8. References / Sources
# List or summarize key URLs or sources used.

# ---

# Below is the content from various sources:
# {full_content}
#         """

#         # Generate report
#         with st.spinner("Generating report with Gemini..."):
#             gemini_response = model.generate_content(prompt)
#             final_report = gemini_response.text

#             # Render report with inline image display
#             for line in final_report.split("\n"):
#                 if line.strip().startswith("[IMAGE]("):
#                     url = line.split("(")[1].split(")")[0]
#                     st.image(url, use_column_width=True)
#                 else:
#                     st.markdown(line)

#             # Export options
#             docx_file = Document()
#             for line in final_report.split("\n"):
#                 docx_file.add_paragraph(line)
#             docx_path = "report.docx"
#             docx_file.save(docx_path)
#             st.download_button("Download as DOCX", open(docx_path, "rb"), file_name="report.docx")

#             pdf = FPDF()
#             pdf.add_page()
#             pdf.set_auto_page_break(auto=True, margin=15)
#             pdf.set_font("Arial", size=12)

#             safe_report = final_report.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"').replace("—", "-")
            
#             def safe_multi_cell(pdf, text, w=0, h=10):
#                 """
#                 Try to print the text with multi_cell.
#                 If it fails due to a long word, split on very long words and retry.
#                 """
#                 try:
#                     pdf.multi_cell(w, h, text)
#                 except Exception as e:
#                     # If error, try splitting on long words
#                     words = text.split(' ')
#                     for word in words:
#                         if len(word) > 50:  # arbitrary long word length threshold
#                             # Insert zero-width space after every 40 chars
#                             word = '\u200b'.join([word[i:i+40] for i in range(0, len(word), 40)])
#                         try:
#                             pdf.multi_cell(w, h, word + ' ')
#                         except Exception:
#                             # fallback: split word forcibly
#                             parts = [word[i:i+40] for i in range(0, len(word), 40)]
#                             for part in parts:
#                                 pdf.multi_cell(w, h, part)

#             for line in safe_report.split("\n"):
#                 safe_multi_cell(pdf, line)

#             pdf_output = pdf.output(dest='S').encode('latin-1', 'replace')  # Replace unsupported chars
#             st.download_button("Download as PDF", data=pdf_output, file_name="report.pdf")



st.set_page_config(page_title="Deep Research Assistant", layout="wide")
st.title("Deep Research Assistant")

mode = st.selectbox("Select Mode", ["Research", "Extract Website Information"])

if mode == "Research":
    st.subheader("Search Filters")
    topic = st.selectbox("Search Topic", ["general", "news", "finance"])

    # Form using updated topic
    with st.form("search_form"):
        col1, col2 = st.columns(2)
        with col1:
            query = st.text_area(label="Query", placeholder="What are the latest updates from Meta?", height=200)
            search_depth = st.selectbox("Search Depth", ["basic", "advanced"])
            max_results = st.slider("Select Depth (max=10)", 1, 10, 5)
            
            
        with col2:
            include_answer = st.selectbox("Include Answer", ["none", "basic", "advanced"])
            include_images = st.checkbox("Include Images")
            country = st.selectbox(
                "Country (only for general)",
                [ "none", "afghanistan", "albania", "algeria", "andorra", "angola", "argentina", "armenia", "australia", "austria", "azerbaijan", "bahamas", "bahrain", "bangladesh", "barbados", "belarus", "belgium", "belize", "benin", "bhutan", "bolivia", "bosnia and herzegovina", "botswana", "brazil", "brunei", "bulgaria", "burkina faso", "burundi", "cambodia", "cameroon", "canada", "cape verde", "central african republic", "chad", "chile", "china", "colombia", "comoros", "congo", "costa rica", "croatia", "cuba", "cyprus", "czech republic", "denmark", "djibouti", "dominican republic", "ecuador", "egypt", "el salvador", "equatorial guinea", "eritrea", "estonia", "ethiopia", "fiji", "finland", "france", "gabon", "gambia", "georgia", "germany", "ghana", "greece", "guatemala", "guinea", "haiti", "honduras", "hungary", "iceland", "india", "indonesia", "iran", "iraq", "ireland", "israel", "italy", "jamaica", "japan", "jordan", "kazakhstan", "kenya", "kuwait", "kyrgyzstan", "latvia", "lebanon", "lesotho", "liberia", "libya", "liechtenstein", "lithuania", "luxembourg", "madagascar", "malawi", "malaysia", "maldives", "mali", "malta", "mauritania", "mauritius", "mexico", "moldova", "monaco", "mongolia", "montenegro", "morocco", "mozambique", "myanmar", "namibia", "nepal", "netherlands", "new zealand", "nicaragua", "niger", "nigeria", "north korea", "north macedonia", "norway", "oman", "pakistan", "panama", "papua new guinea", "paraguay", "peru", "philippines", "poland", "portugal", "qatar", "romania", "russia", "rwanda", "saudi arabia", "senegal", "serbia", "singapore", "slovakia", "slovenia", "somalia", "south africa", "south korea", "south sudan", "spain", "sri lanka", "sudan", "sweden", "switzerland", "syria", "taiwan", "tajikistan", "tanzania", "thailand", "togo", "trinidad and tobago", "tunisia", "turkey", "turkmenistan", "uganda", "ukraine", "united arab emirates", "united kingdom", "united states", "uruguay", "uzbekistan", "venezuela", "vietnam", "yemen", "zambia", "zimbabwe"],
                disabled=(topic != "general")
            )
            days = st.slider("Days (only for news)", 1, 7, 6, disabled=(topic != "news"))
            time_range = st.selectbox("Time Range", ["none", "day", "week", "month", "year"])

        submitted = st.form_submit_button("Search")

    if submitted:
        with st.spinner("Summoning the wisdom of the web"):
            search_params = {
                "query": query,
                "search_depth": search_depth,
                "max_results": max_results,
                "time_range": time_range,
                "include_answer": include_answer,
                "include_images": include_images,
               
            }
            if time_range == "none":
                del search_params["time_range"]
            if include_answer == "none":
                search_params["include_answer"] = False
            if topic == "news":
                search_params["days"] = days
                if "country" in search_params:
                    del search_params["country"]
            if topic == "general":
                search_params["country"] = country
            if country == "none":
                if "country" in search_params:
                    del search_params["country"]
            
            response = client.search(**search_params)
            if len(response["results"]) < 1:
                st.error("No results found!")
                
            else:
                st.success("Search completed!")
                st.json(response)
            
                content_blocks = []
                for result in response["results"]:
                    text = result.get("content", "")
                    images = result.get("images", [])
                    content_blocks.append(text)
                    for img_url in images:
                        content_blocks.append(f"[IMAGE]({img_url})")

                full_content = "\n\n".join(content_blocks)

                prompt = f"""
    You are an expert research analyst and report writer. Your task is to create a deep, detailed, structured, and insightful research report on the following topic:

    **Topic:** {query}

    Use the provided content below as your primary source of information. Analyze it thoroughly. Extract facts, insights, and diverse perspectives. Your report should follow this structure:

    ---

    ### 1. Executive Summary
    Summarize the key findings and importance of the topic in under 150 words.

    ### 2. Introduction
    Explain the background and context. Why is it relevant now?

    ### 3. Key Highlights
    - Summarize major findings or trends.
    - Use bullet points.

    ### 4. In-depth Analysis
    Break down the topic into multiple sections. For each:
    - Explain clearly.
    - Use quotes or paraphrased points from sources.
    - Mention supporting images where relevant.

    ### 5. Recent Developments (if any)
    Highlight newsworthy events.

    ### 6. Regional or Industry Insights
    (Optional based on content)

    ### 7. Conclusion
    Summarize and suggest future outlooks or actions.

    ### 8. References / Sources
    List or summarize key URLs or sources used.

    ---

    Below is the content from various sources:
    {full_content}
    """

                with st.spinner("Generating response..."):
                    gemini_response = model.generate_content(prompt)
                    st.session_state.final_report = gemini_response.text

                    for line in st.session_state.get("final_report", "").split("\n"):
                        if line.strip().startswith("[IMAGE]("):
                            url = line.split("(")[1].split(")")[0]
                            st.image(url, use_column_width=True)
                        else:
                            st.markdown(line)

                    docx_file = Document()
                    for line in st.session_state.get("final_report", "").split("\n"):
                        docx_file.add_paragraph(line)

                    docx_buffer = BytesIO()
                    docx_file.save(docx_buffer)
                    docx_buffer.seek(0)

                    st.download_button(
                        label="Download as DOCX",
                        data=docx_buffer,
                        file_name="report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

elif mode == "Extract Website Information":
    st.subheader("Website Content Extractor")
    url_input = st.text_input("Enter website URL to extract information")
    extract_query = st.text_area("What do you want to extract from this website? (e.g., summarize, find trends, get company info)", height=200)
    extract_btn = st.button("Extract and Generate Report")

    if extract_btn and url_input:
        with st.spinner("Interrogating this website..."):
            extract_response = client.extract(
                urls=[url_input],
                extract_depth="advanced",
                include_images=True
            )
            st.success("Extraction completed!")
            st.json(extract_response)
            extract_content_blocks = []
            for doc in extract_response["results"]:
                text = doc.get("raw_content", "")
                images = doc.get("images", [])
                extract_content_blocks.append(text)
                for img_url in images:
                    extract_content_blocks.append(f"[IMAGE]({img_url})")

            full_extract = "\n\n".join(extract_content_blocks)

            extract_prompt = f"""
You are a web research assistant. Based on the extracted website content below, please fulfill the following request:

### Request:
{extract_query}

---

### Website Content:
{full_extract}

---

Generate a clean, structured markdown report that fulfills the user request above. Include image references where appropriate.
"""

            with st.spinner("Generating response..."):
                gemini_response = model.generate_content(extract_prompt)
                st.session_state.extract_report = gemini_response.text

                for line in st.session_state.get("extract_report", "").split("\n"):
                    if line.strip().startswith("[IMAGE]("):
                        url = line.split("(")[1].split(")")[0]
                        st.image(url, use_column_width=True)
                    else:
                        st.markdown(line)

                docx_file = Document()
                for line in st.session_state.get("extract_report", "").split("\n"):
                    docx_file.add_paragraph(line)

                docx_buffer = BytesIO()
                docx_file.save(docx_buffer)
                docx_buffer.seek(0)

                st.download_button(
                    label="Download as DOCX",
                    data=docx_buffer,
                    file_name="extract_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
